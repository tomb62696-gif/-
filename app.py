import streamlit as st
import docx
from docx import Document
import openai
import io

# 1. 网页基础配置
st.set_page_config(page_title="地铁车站结构设计审查评价系统", layout="wide")
st.title("🚇 地铁车站结构设计审查评价系统 (AI Agent)")
st.caption("基于国家 5 大核心规范 —— 强条硬核管控版（支持报告导出）")

# 2. 侧边栏配置
with st.sidebar:
    st.header("⚙️ 审查配置中心")
    api_key = st.text_input("请输入大模型 API Key", type="password")
    base_url = st.text_input("API 端点 URL", value="https://api.deepseek.com/v1")
    model_name = st.text_input("大模型名称", value="deepseek-chat")
    
    st.divider()
    st.subheader("📚 核心校审规范依据")
    st.markdown("""
    1. **建质[2013]160号** 编制深度规定
    2. **发改基础[2015]49号** 规划建设管理通知
    3. **建标104-2008** 工程项目建设标准
    4. **GB50490-2009** 城市轨道交通技术规范
    5. **GB50157-2013** 地铁设计规范
    """)

# 3. Word 文档输入解析器
def extract_word_content(file):
    doc = docx.Document(file)
    content_lines = []
    for para in doc.paragraphs:
        if para.text.strip(): content_lines.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_data: content_lines.append(f"[表格] " + " | ".join(row_data))
    return "\n".join(content_lines)

# ⭐ 新增：将 AI 的 Markdown 报告转换为 Word 二进制流的函数
def convert_md_to_docx(md_text, report_title="地铁车站结构设计审查报告"):
    doc = Document()
    
    # 添加大标题
    doc.add_heading(report_title, level=0)
    
    # 简单的 Markdown 行解析器
    lines = md_text.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # 解析各级标题
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        # 解析无序列表
        elif line.startswith('* ') or line.startswith('- '):
            # 去除可能存在的 Markdown 加粗符号 **
            clean_item = line[2:].replace('**', '')
            doc.add_paragraph(clean_item, style='List Bullet')
        # 解析普通正文
        else:
            clean_text = line.replace('**', '') # 清理加粗符号
            doc.add_paragraph(clean_text)
            
    # 将文件保存到内存缓冲区（避免在服务器留下临时文件）
    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    return docx_buffer.getvalue()


# 4. 文件上传
uploaded_file = st.file_uploader("📂 请上传地铁车站设计文件 Word 文档 (.docx)", type=["docx"])

if uploaded_file:
    with st.spinner("正在解析上传的 Word 文档..."):
        document_text = extract_word_content(uploaded_file)
    st.success(f"🎉 文档解析成功！共提取 {len(document_text)} 个字符。")

    # 5. 开始审查
    if st.button("🚀 开始自动化规范强条审查", type="primary"):
        if not api_key:
            st.error("❌ 请先在左侧边栏配置您的 API Key。")
        else:
            with st.spinner("🕵️ AI 强条猎手正在全文本扫描违规风险..."):
                try:
                    system_prompt = """你是一个地铁车站结构设计强条审查官。请严格审查文本。
开篇必须立刻呈现【🚨强条合规判定结果】，如果有违规，必须列出【🛑 强条违规/漏项清单】置顶。
每项违规必须满足：[文档原文定位] + [违反的规范条文号] + [整改要求]。"""

                    client = openai.OpenAI(api_key=api_key, base_url=base_url)
                    response = client.chat.completions.create(
                        model=model_name,
                        messages=[
                            {"role": "system", "content": system_prompt},
                            {"role": "user", "content": f"请审查以下内容：\n\n{document_text}"}
                        ],
                        temperature=0.0
                    )
                    
                    report_content = response.choices[0].message.content
                    
                    # 将报告内容存入 Streamlit 的 SessionState 缓存中，防止页面刷新丢失
                    st.session_state['report_content'] = report_content
                    
                except Exception as e:
                    st.error(f"审查引擎运行中断: {str(e)}")

    # 6. ⭐ 渲染报告与下载组件
    if 'report_content' in st.session_state:
        st.divider()
        st.header("📊 AI 自动化审查评价报告")
        
        # 风险置顶状态框
        if "违反" in st.session_state['report_content'] or "违规" in st.session_state['report_content']:
            st.error("⚠️ 系统检测到当前设计文件存在违反强制性条文的风险，已自动触发风险阻断！")
        else:
            st.success("✅ 初步强条扫描未见明显违规。")
            
        # 展现报告文本
        st.markdown(st.session_state['report_content'])
        
        st.divider()
        st.subheader("💾 导出与下载报告")
        
        # 列布局：并排下载
        col1, col2 = st.columns(2)
        
        with col1:
            # 生成 Word 二进制流
            with st.spinner("正在打包 Word 文件..."):
                docx_bytes = convert_md_to_docx(st.session_state['report_content'])
            
            # Streamlit 原生下载按钮
            st.download_button(
                label="📥 下载 Word (.docx) 格式报告",
                data=docx_bytes,
                file_name="地铁车站结构设计审查报告.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
            
        with col2:
            st.info("💡 **提示（导出 PDF）：** 键盘按下 `Ctrl + P` (Mac用户按 `Cmd + P`)，在弹出的系统打印窗口中将目标打印机选择为 **“另存为 PDF”**，即可完美导出包含网页高亮色彩的 PDF 报告。")
